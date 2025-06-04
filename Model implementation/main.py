import numpy as np
import matplotlib.pyplot as plt
from matplotlib.animation import FuncAnimation

def get_impulse(t):
    phase = int(t // impulse_period) % 2
    in_impulse = (t % impulse_period) < impulse_duration
    if not in_impulse:
        return 0.0
    return impulse_strength_left if phase == 0 else impulse_strength_right

def f(current_time, X):
    theta1, theta2, omega1, omega2 = X
    delta = theta2 - theta1

    denominator1 = segment_length * (2 - np.cos(2 * delta))
    denominator2 = denominator1

    impulse = get_impulse(current_time)

    theta_derivative = omega1
    theta2_derivative = omega2

    num1 = (-gravity * (2 * np.sin(theta1)) - np.sin(delta) * (omega2 ** 2 * segment_length + omega1 ** 2 * segment_length * np.cos(delta))
            + impulse)
    omega1_derivative = (num1 - damping * omega1) / denominator1

    num2 = (2 * np.sin(delta) * (omega1 ** 2 * segment_length + gravity * np.cos(theta1)))
    omega2_derivative = (num2 - damping * omega2) / denominator2

    return np.array([theta_derivative, theta2_derivative, omega1_derivative, omega2_derivative])

def rk4_method(f, t, X, time_step):
    k1 = f(t, X)
    k2 = f(t + time_step / 2, X + time_step / 2 * k1)
    k3 = f(t + time_step / 2, X + time_step / 2 * k2)
    k4 = f(t + time_step, X + time_step * k3)
    return X + time_step / 6 * (k1 + 2 * k2 + 2 * k3 + k4)

def add_impulse_spans(ax, impulse_times, color, label):
    for t_imp in impulse_times:
        ax.axvspan(t_imp, t_imp + impulse_duration, color=color, alpha=0.1,
                   label=label if t_imp == impulse_times[0] else "")

def run_simulation():
    for i in range(1, len(times)):
        t = times[i - 1]
        impulse = get_impulse(t)
        if impulse == impulse_strength_left:
            impulse_times_left.append(t)
        elif impulse == impulse_strength_right:
            impulse_times_right.append(t)
        X[i] = rk4_method(f, t, X[i - 1], time_step)

def create_plots(generate_report=False):
    fig_plots, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 8))

    ax1.plot(times, X[:, 0], label='Kąt')
    ax1.plot(times, X[:, 2], label='Prędkość kątowa')
    add_impulse_spans(ax1, impulse_times_left, 'blue', 'Impuls lewy')
    add_impulse_spans(ax1, impulse_times_right, 'green', 'Impuls prawy')
    ax1.set_ylabel('Wartość')
    ax1.set_title('Pierwsza część wahadła')
    ax1.legend(loc='upper left')
    ax1.grid(True)

    ax2.plot(times, X[:, 1], label='Kąt')
    ax2.plot(times, X[:, 3], label='Prędkość kątowa')
    add_impulse_spans(ax2, impulse_times_left, 'blue', 'Impuls lewy')
    add_impulse_spans(ax2, impulse_times_right, 'green', 'Impuls prawy')
    ax2.set_xlabel('Czas [s]')
    ax2.set_ylabel('Wartość')
    ax2.set_title('Druga część wahadła')
    ax2.legend(loc='upper left')
    ax2.grid(True)

    plt.tight_layout()

    fig_animation, ax = plt.subplots(figsize=(6, 6))
    ax.set_xlim(-5, 5)
    ax.set_ylim(-5, 5)

    pendulum_line = ax.plot([], [], 'o-', lw=3, color='tab:red')[0]
    time_text = ax.text(0.05, 0.95, '', transform=ax.transAxes)

    def init_func():
        pendulum_line.set_data([], [])
        time_text.set_text('')
        return pendulum_line, time_text, impulse_left_rect, impulse_right_rect

    def update_func(frame):
        theta1, theta2 = X[frame, 0], X[frame, 1]
        x1 = segment_length * np.sin(theta1)
        y1 = -segment_length * np.cos(theta1)
        x2 = x1 + segment_length * np.sin(theta2)
        y2 = y1 - segment_length * np.cos(theta2)
        pendulum_line.set_data([0, x1, x2], [0, y1, y2])
        time_text.set_text(f'Czas = {times[frame]:.2f} s')

        t = times[frame]
        in_impulse = (t % impulse_period) < impulse_duration
        phase = int(t // impulse_period) % 2
        impulse_left_rect.set_color('blue' if in_impulse and phase == 0 else 'black')
        impulse_right_rect.set_color('green' if in_impulse and phase == 1 else 'black')

        return pendulum_line, time_text, impulse_left_rect, impulse_right_rect

    impulse_left_rect = plt.Rectangle((-3.25, -1), 0.5, 0.5, color='black', alpha=0.5)
    impulse_right_rect = plt.Rectangle((3.25, -1), 0.5, 0.5, color='black', alpha=0.5)

    ax.add_patch(impulse_left_rect)
    ax.add_patch(impulse_right_rect)


    _ = FuncAnimation(fig_animation, update_func, frames=len(times), init_func=init_func,
                        blit=True, interval=10)

    if generate_report:
        return fig_plots
    plt.show()

def run_test():
    run_simulation()
    create_plots(generate_report=False)

def generate_report():
    from docx import Document
    from docx2pdf import convert
    from os import remove
    from io import BytesIO
    from docx.shared import Inches

    document = Document()
    document.add_heading('Zadanie domowe - implementacja modelu\nWojciech Latos', level=1)
    document.add_paragraph('Wykonano symulację modelu dwuczęściowego wahadła. Ruch odbywa się na podstawie siły impulsów, które są naprzemienne generowane raz z prawej raz z lewej strony. '
                           'W momencie aktywacji impulsu, kwadrat który go reprezentuje na animacji zmienia kolor.'
                           'Siłę impulsów, długość trwania impulsu, okres co jaki czas impuls zostanie wykonany, długość wahadła, moc tłumienia oraz opcjonalnie grawitację można zmieniać według własnych preferencji. '
                           'Wykorzystano metodę Runge-Kutta 4. rzędu do numerycznego rozwiązania układu równań ruchu. '
                           'Na poniższych wykresach przedstawiono przebieg zmian kątów i prędkości kątowych obu części wahadeł oraz momenty, w których dane źródła energii były aktywne.\n'
                           'Dodatkowo zostały dodane kilka zrzutów ekranów z okna wizualizacji modelu. Uruchamiając kod można zobaczyć animację w czasie rzeczywistym.')

    run_simulation()

    fig = create_plots(generate_report=True)

    memfile = BytesIO()
    fig.savefig(memfile)
    plt.close(fig)
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(7))
    memfile.close()

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

if __name__ == "__main__":
    gravity = 9.81
    segment_length = 1.5
    damping = 0.1

    impulse_strength_left = 20.0
    impulse_strength_right = 10.0
    impulse_period = 3
    impulse_duration = 0.2

    t0, t_end = 0, 20
    time_step = 0.01
    times = np.arange(t0, t_end, time_step)
    X = np.zeros((len(times), 4))
    X[0] = [0, 0, 0, 0]

    impulse_times_left = []
    impulse_times_right = []

    run_test()
    # generate_report()
